import re
import mimetypes
import magic
import structlog
from typing import Dict, List, Optional
from pathlib import Path

# Import document processing libraries
try:
    import PyPDF2
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

logger = structlog.get_logger()


def get_file_type(file_path: str) -> str:
    """Determine file MIME type"""
    try:
        # Use python-magic for accurate detection
        mime = magic.Magic(mime=True)
        return mime.from_file(file_path)
    except Exception:
        # Fallback to mimetypes
        mime_type, _ = mimetypes.guess_type(file_path)
        return mime_type or "application/octet-stream"


def extract_text_from_file(file_path: str, file_extension: str) -> str:
    """Extract text content from various file types"""
    try:
        file_extension = file_extension.lower()
        
        if file_extension == '.txt':
            return extract_text_from_txt(file_path)
        elif file_extension == '.md':
            return extract_text_from_markdown(file_path)
        elif file_extension == '.pdf':
            return extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return extract_text_from_docx(file_path)
        elif file_extension in ['.xlsx', '.xls']:
            return extract_text_from_xlsx(file_path)
        elif file_extension == '.csv':
            return extract_text_from_csv(file_path)
        else:
            logger.warning("Unsupported file type for text extraction", extension=file_extension)
            return ""
            
    except Exception as e:
        logger.error("Text extraction failed", file_path=file_path, error=str(e))
        return ""


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from plain text files"""
    encodings = ['utf-8', 'utf-16', 'iso-8859-1', 'cp1252']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                return file.read()
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.error("Failed to read text file", file_path=file_path, error=str(e))
            return ""
    
    logger.error("Could not decode text file with any encoding", file_path=file_path)
    return ""


def extract_text_from_markdown(file_path: str) -> str:
    """Extract text from Markdown files (basic implementation)"""
    try:
        text = extract_text_from_txt(file_path)
        
        # Remove Markdown formatting (basic)
        # Remove headers
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
        
        # Remove bold and italic
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'__(.*?)__', r'\1', text)
        text = re.sub(r'_(.*?)_', r'\1', text)
        
        # Remove links
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
        
        # Remove code blocks
        text = re.sub(r'```[\s\S]*?```', '', text)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        
        # Remove horizontal rules
        text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)
        
        return text.strip()
        
    except Exception as e:
        logger.error("Markdown text extraction failed", file_path=file_path, error=str(e))
        return ""


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF files"""
    if not PDF_AVAILABLE:
        logger.error("PyPDF2 not available for PDF text extraction")
        return ""
    
    try:
        text_content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                except Exception as e:
                    logger.warning("Failed to extract text from PDF page", 
                                 page_num=page_num, error=str(e))
                    continue
        
        return '\n\n'.join(text_content)
        
    except Exception as e:
        logger.error("PDF text extraction failed", file_path=file_path, error=str(e))
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX files"""
    if not DOCX_AVAILABLE:
        logger.error("python-docx not available for DOCX text extraction")
        return ""
    
    try:
        doc = DocxDocument(file_path)
        text_content = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_text.append(' | '.join(row_text))
            
            if table_text:
                text_content.append('\n'.join(table_text))
        
        return '\n\n'.join(text_content)
        
    except Exception as e:
        logger.error("DOCX text extraction failed", file_path=file_path, error=str(e))
        return ""


def extract_text_from_xlsx(file_path: str) -> str:
    """Extract text from Excel files"""
    if not XLSX_AVAILABLE:
        logger.error("openpyxl not available for XLSX text extraction")
        return ""
    
    try:
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        text_content = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_text = [f"Sheet: {sheet_name}"]
            
            for row in sheet.iter_rows(values_only=True):
                row_text = []
                for cell in row:
                    if cell is not None and str(cell).strip():
                        row_text.append(str(cell).strip())
                
                if row_text:
                    sheet_text.append(' | '.join(row_text))
            
            if len(sheet_text) > 1:  # More than just the sheet name
                text_content.append('\n'.join(sheet_text))
        
        return '\n\n'.join(text_content)
        
    except Exception as e:
        logger.error("XLSX text extraction failed", file_path=file_path, error=str(e))
        return ""


def extract_text_from_csv(file_path: str) -> str:
    """Extract text from CSV files"""
    try:
        import csv
        text_content = []
        
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'iso-8859-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, newline='') as file:
                    # Try to detect delimiter
                    sample = file.read(1024)
                    file.seek(0)
                    
                    sniffer = csv.Sniffer()
                    delimiter = sniffer.sniff(sample).delimiter
                    
                    reader = csv.reader(file, delimiter=delimiter)
                    
                    for row_num, row in enumerate(reader):
                        if row_num > 1000:  # Limit to first 1000 rows
                            break
                        
                        row_text = ' | '.join([str(cell).strip() for cell in row if str(cell).strip()])
                        if row_text:
                            text_content.append(row_text)
                    
                    break  # Successfully processed with this encoding
                    
            except (UnicodeDecodeError, csv.Error):
                continue
            except Exception as e:
                logger.warning("CSV processing error with encoding", 
                             encoding=encoding, error=str(e))
                continue
        
        return '\n'.join(text_content)
        
    except Exception as e:
        logger.error("CSV text extraction failed", file_path=file_path, error=str(e))
        return ""


def clean_text(text: str) -> str:
    """Clean and normalize extracted text"""
    if not text:
        return ""
    
    try:
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive newlines
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        
        # Remove non-printable characters except newlines and tabs
        text = re.sub(r'[^\x20-\x7E\n\t]', '', text)
        
        # Remove excessive punctuation
        text = re.sub(r'([.!?]){2,}', r'\1', text)
        
        # Remove URLs
        text = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '', text)
        
        # Remove email addresses
        text = re.sub(r'\S+@\S+\.\S+', '', text)
        
        return text.strip()
        
    except Exception as e:
        logger.error("Text cleaning failed", error=str(e))
        return text


def extract_keywords(text: str, max_keywords: int = 20) -> List[str]:
    """Extract keywords from text (simple implementation)"""
    try:
        if not text:
            return []
        
        # Simple keyword extraction using frequency
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        
        # Remove common stop words
        stop_words = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with',
            'by', 'from', 'up', 'about', 'into', 'through', 'during', 'before', 'after', 'above',
            'below', 'under', 'between', 'among', 'this', 'that', 'these', 'those', 'i', 'me',
            'my', 'myself', 'we', 'our', 'ours', 'ourselves', 'you', 'your', 'yours',
            'yourself', 'yourselves', 'he', 'him', 'his', 'himself', 'she', 'her', 'hers',
            'herself', 'it', 'its', 'itself', 'they', 'them', 'their', 'theirs', 'themselves',
            'what', 'which', 'who', 'whom', 'whose', 'where', 'when', 'why', 'how', 'all',
            'any', 'both', 'each', 'few', 'more', 'most', 'other', 'some', 'such', 'no',
            'nor', 'not', 'only', 'own', 'same', 'so', 'than', 'too', 'very', 'can', 'will',
            'just', 'should', 'now', 'are', 'was', 'were', 'been', 'being', 'have', 'has',
            'had', 'having', 'do', 'does', 'did', 'doing', 'would', 'could', 'should',
            'may', 'might', 'must', 'shall', 'will', 'am', 'is'
        }
        
        # Filter out stop words and count frequency
        word_freq = {}
        for word in words:
            if word not in stop_words and len(word) > 3:
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Sort by frequency and return top keywords
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        keywords = [word for word, freq in sorted_words[:max_keywords] if freq > 1]
        
        return keywords
        
    except Exception as e:
        logger.error("Keyword extraction failed", error=str(e))
        return []


def extract_entities(text: str, max_entities: int = 50) -> List[str]:
    """Extract named entities from text (simple implementation)"""
    try:
        if not text:
            return []
        
        entities = []
        
        # Extract capitalized words/phrases (potential entities)
        # This is a simple implementation - in production you'd use NLP libraries like spaCy
        capitalized_pattern = r'\b[A-Z][a-zA-Z]*(?:\s+[A-Z][a-zA-Z]*)*\b'
        matches = re.findall(capitalized_pattern, text)
        
        # Filter and clean entities
        for match in matches:
            entity = match.strip()
            if (len(entity) > 2 and 
                entity not in ['The', 'This', 'That', 'These', 'Those', 'And', 'Or', 'But'] and
                not entity.isupper()):  # Avoid all-caps words
                entities.append(entity)
        
        # Remove duplicates while preserving order
        seen = set()
        unique_entities = []
        for entity in entities:
            if entity.lower() not in seen:
                seen.add(entity.lower())
                unique_entities.append(entity)
        
        return unique_entities[:max_entities]
        
    except Exception as e:
        logger.error("Entity extraction failed", error=str(e))
        return []


def summarize_text(text: str, max_sentences: int = 5) -> str:
    """Create a simple extractive summary of text"""
    try:
        if not text:
            return ""
        
        # Split into sentences
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 10]
        
        if len(sentences) <= max_sentences:
            return '. '.join(sentences) + '.'
        
        # Simple scoring based on word frequency and position
        word_freq = {}
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        
        for word in words:
            word_freq[word] = word_freq.get(word, 0) + 1
        
        # Score sentences
        sentence_scores = []
        for i, sentence in enumerate(sentences):
            score = 0
            words_in_sentence = re.findall(r'\b[a-zA-Z]{3,}\b', sentence.lower())
            
            # Frequency score
            for word in words_in_sentence:
                score += word_freq.get(word, 0)
            
            # Position score (first and last sentences are often important)
            if i < 2:  # First two sentences
                score *= 1.5
            elif i >= len(sentences) - 2:  # Last two sentences
                score *= 1.2
            
            # Length penalty for very short sentences
            if len(words_in_sentence) < 5:
                score *= 0.5
            
            sentence_scores.append((score, sentence))
        
        # Select top sentences
        top_sentences = sorted(sentence_scores, key=lambda x: x[0], reverse=True)[:max_sentences]
        
        # Sort by original order
        selected_sentences = []
        for score, sentence in top_sentences:
            original_index = sentences.index(sentence)
            selected_sentences.append((original_index, sentence))
        
        selected_sentences.sort(key=lambda x: x[0])
        summary = '. '.join([sentence for _, sentence in selected_sentences]) + '.'
        
        return summary
        
    except Exception as e:
        logger.error("Text summarization failed", error=str(e))
        return text[:500] + "..." if len(text) > 500 else text


def detect_language(text: str) -> str:
    """Detect language of text (simple implementation)"""
    try:
        if not text:
            return "unknown"
        
        # Simple language detection based on common words
        # This is very basic - in production use proper language detection libraries
        
        sample = text.lower()[:1000]  # Use first 1000 characters
        
        # English indicators
        english_words = ['the', 'and', 'is', 'in', 'to', 'of', 'a', 'that', 'it', 'with', 'for', 'as', 'was', 'on', 'are']
        english_score = sum(1 for word in english_words if word in sample)
        
        # Spanish indicators
        spanish_words = ['el', 'la', 'de', 'que', 'y', 'en', 'un', 'es', 'se', 'no', 'te', 'lo', 'le', 'da', 'su']
        spanish_score = sum(1 for word in spanish_words if word in sample)
        
        # French indicators
        french_words = ['le', 'de', 'et', 'à', 'un', 'il', 'être', 'et', 'en', 'avoir', 'que', 'pour', 'dans', 'ce', 'son']
        french_score = sum(1 for word in french_words if word in sample)
        
        scores = {
            'english': english_score,
            'spanish': spanish_score,
            'french': french_score
        }
        
        if max(scores.values()) > 0:
            return max(scores, key=scores.get)
        else:
            return "unknown"
            
    except Exception as e:
        logger.error("Language detection failed", error=str(e))
        return "unknown"


def validate_file_content(file_path: str, expected_type: str = None) -> Dict[str, any]:
    """Validate file content and return metadata"""
    try:
        file_info = {
            "is_valid": False,
            "file_size": 0,
            "mime_type": "",
            "encoding": "",
            "has_text_content": False,
            "estimated_pages": 0,
            "word_count": 0,
            "character_count": 0,
            "errors": []
        }
        
        # Check if file exists
        if not Path(file_path).exists():
            file_info["errors"].append("File does not exist")
            return file_info
        
        # Get file size
        file_info["file_size"] = Path(file_path).stat().st_size
        
        # Get MIME type
        file_info["mime_type"] = get_file_type(file_path)
        
        # Try to extract text to validate content
        file_extension = Path(file_path).suffix.lower()
        text_content = extract_text_from_file(file_path, file_extension)
        
        if text_content:
            file_info["has_text_content"] = True
            file_info["word_count"] = len(text_content.split())
            file_info["character_count"] = len(text_content)
            file_info["estimated_pages"] = max(1, file_info["word_count"] // 250)  # ~250 words per page
            file_info["is_valid"] = True
        else:
            file_info["errors"].append("No text content could be extracted")
        
        # Validate against expected type if provided
        if expected_type and expected_type not in file_info["mime_type"]:
            file_info["errors"].append(f"File type mismatch. Expected: {expected_type}, Got: {file_info['mime_type']}")
        
        return file_info
        
    except Exception as e:
        logger.error("File validation failed", file_path=file_path, error=str(e))
        return {
            "is_valid": False,
            "errors": [f"Validation error: {str(e)}"]
        }


def chunk_text_intelligently(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, any]]:
    """Split text into intelligent chunks preserving semantic boundaries"""
    try:
        if not text:
            return []
        
        chunks = []
        
        # Split by paragraphs first
        paragraphs = re.split(r'\n\s*\n', text)
        
        current_chunk = ""
        current_size = 0
        chunk_index = 0
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            paragraph_size = len(paragraph)
            
            # If single paragraph is larger than chunk_size, split by sentences
            if paragraph_size > chunk_size:
                if current_chunk:
                    # Save current chunk
                    chunks.append({
                        "index": chunk_index,
                        "content": current_chunk.strip(),
                        "size": len(current_chunk),
                        "type": "paragraph_boundary"
                    })
                    chunk_index += 1
                    current_chunk = ""
                    current_size = 0
                
                # Split large paragraph by sentences
                sentences = re.split(r'[.!?]+', paragraph)
                sentence_chunk = ""
                
                for sentence in sentences:
                    sentence = sentence.strip()
                    if not sentence:
                        continue
                    
                    if len(sentence_chunk) + len(sentence) > chunk_size:
                        if sentence_chunk:
                            chunks.append({
                                "index": chunk_index,
                                "content": sentence_chunk.strip(),
                                "size": len(sentence_chunk),
                                "type": "sentence_boundary"
                            })
                            chunk_index += 1
                            
                            # Add overlap from previous chunk
                            if overlap > 0 and len(sentence_chunk) > overlap:
                                sentence_chunk = sentence_chunk[-overlap:] + " " + sentence
                            else:
                                sentence_chunk = sentence
                        else:
                            sentence_chunk = sentence
                    else:
                        sentence_chunk += " " + sentence if sentence_chunk else sentence
                
                if sentence_chunk:
                    current_chunk = sentence_chunk
                    current_size = len(sentence_chunk)
            
            # Normal paragraph processing
            elif current_size + paragraph_size > chunk_size:
                if current_chunk:
                    chunks.append({
                        "index": chunk_index,
                        "content": current_chunk.strip(),
                        "size": len(current_chunk),
                        "type": "paragraph_boundary"
                    })
                    chunk_index += 1
                    
                    # Add overlap
                    if overlap > 0 and len(current_chunk) > overlap:
                        current_chunk = current_chunk[-overlap:] + "\n\n" + paragraph
                    else:
                        current_chunk = paragraph
                    current_size = len(current_chunk)
                else:
                    current_chunk = paragraph
                    current_size = paragraph_size
            else:
                current_chunk += "\n\n" + paragraph if current_chunk else paragraph
                current_size += paragraph_size + (2 if current_chunk else 0)
        
        # Add final chunk
        if current_chunk:
            chunks.append({
                "index": chunk_index,
                "content": current_chunk.strip(),
                "size": len(current_chunk),
                "type": "final"
            })
        
        return chunks
        
    except Exception as e:
        logger.error("Intelligent chunking failed", error=str(e))
        # Fallback to simple chunking
        return simple_chunk_text(text, chunk_size, overlap)


def simple_chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, any]]:
    """Simple text chunking as fallback"""
    try:
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at word boundary
            if end < len(text):
                # Look back for space
                while end > start and text[end] != ' ':
                    end -= 1
                
                if end == start:  # No space found, force break
                    end = start + chunk_size
            
            chunk_content = text[start:end].strip()
            
            if chunk_content:
                chunks.append({
                    "index": chunk_index,
                    "content": chunk_content,
                    "size": len(chunk_content),
                    "type": "simple"
                })
                chunk_index += 1
            
            start = max(start + 1, end - overlap)
        
        return chunks
        
    except Exception as e:
        logger.error("Simple chunking failed", error=str(e))
        return [] 